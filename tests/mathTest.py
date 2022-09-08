import docx
#also need python-docx package installed for docx
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


class ArithmeticQuestion:
    operator_translation = {
        '+': '+',
        '-': '-',
        '*': '\u00D7',
        '/': '\u00F7'
    }

    def __init__(self, operand_1, operator_1, operand_2, operator_2=None, operand_3=None, operator_3=None, operand_4=None):
        self.operand_1 = operand_1
        self.operand_2 = operand_2
        self.operand_3 = operand_3
        self.operand_4 = operand_4
        self.operator = operator_1
        self.operator = operator_2
        self.operator = operator_3

        if not (operator_2 or operator_3):
            self.answer = eval(str(operand_1)+operator_1+str(operand_2))
            self.num_operands = 2

        elif not (operator_3):
            self.answer = eval(str(operand_1)+operator_1+str(operand_2)+operator_2+str(operand_3))
            self.num_operands = 3

        else:
            self.answer = eval(str(operand_1) + operator_1 + str(operand_2) + operator_2 + str(operand_3) + operator_3 + str(operand_4))
            self.num_operands = 4


if __name__ == "__main__":
    now = datetime.now().strftime("%m_%d_%Y %H_%M_%S")

    questions_path = os.path.join(os.getcwd(), 'questions')
    answers_path = os.path.join(os.getcwd(), 'answers')
    if not os.path.exists(questions_path):
        print('Making questions folder: ', questions_path)
        os.makedirs(questions_path)

    if not os.path.exists(answers_path):
        print('Making answers folder: ', answers_path)
        os.makedirs(answers_path)

    test_name = 'Math 7 Unit 1 Computation with Integers Show Your Learning.docx'
    question_file_path = questions_path + '/' + now + ' questions ' + test_name
    answer_file_path = answers_path + '/' + now + ' answers ' + test_name
    num_rows = 16
    num_cols = 6
    document = docx.Document()
    document.add_heading('Math 7 Show Your Learning', 1)
    document.add_heading('Computation with Integers			Class:			Name:', 2)

    table = document.add_table(rows=num_rows, cols=num_cols, style='Table Grid')

    column_widths = {
        0: 2.72,
        1: 3.73,
        2: 5.29,
        3: 7.2,
        4: 1.16,
        5: 1.48
    }

    for k, v in column_widths.items():
        for cell in table.columns[k].cells:
            cell.width = Cm(v)

    column_headers = {
        0: 'Computation\nwith integers',
        1: '2\nBasic',
        2: '3\nIntermediate',
        3: '4\nAdvanced',
        4: 'Out\nof',
        5: 'My\nMark'
    }

    for k, v in column_headers.items():
        table.cell(0, k).text = v
        table.cell(0, k).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row_headers = {
        1: 'Multiplication\nTo 999x999',
        3: 'Division\nUnder\n1,000,000',
        5: 'Adding \u00B1\nintegers',
        7: 'Subtracting\n\u00B1 integers',
        9: 'Multiplying\n\u00B1 integers',
        11: 'Dividing\n\u00B1 integers',
        13: 'Order of\nOperations\nwith integers',
        15: 'Total',
    }

    for k, v in row_headers.items():
        table.cell(k, 0).text = v

    column_out_of = {
        1: 4,
        3: 4,
        5: 4,
        7: 4,
        9: 4,
        11: 4,
        13: 4,
        15: 28,
    }

    out_of_column = 4
    for k, v in column_out_of.items():
        table.cell(k, out_of_column).text = str(v)
        table.cell(k, out_of_column).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for x in range(1, 4):
        table.cell(15, x).text = 'N/A'
        table.cell(15, x).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading('Instructions:', 2)
    p = document.add_paragraph('Answer ', style='List Bullet')
    p.add_run('only ').bold = True
    p.add_run('the questions you haven’t yet demonstrated your level of mastery for.')
    document.add_paragraph('If you show mastery of a skill at a certain level twice IN A ROW, you get credit.You can do'
                           ' that on this page, or you might have shown it 1 or 2 times elsewhere. Check your own '
                           'learning map.', style='List Bullet')
    document.add_paragraph('If you master “Advanced”, or “Intermediate”, you automatically get credit for the easier'
                           ' levels.', style='List Bullet')
    document.add_paragraph('You will likely need more space. Use the blank side of the paper, label your questions, and'
                           ' show your steps.', style='List Bullet')

    document.save(question_file_path)
    print('Questions outputted to file: ', question_file_path)

    document.save(answer_file_path)
    print('Answers outputted to file: ', answer_file_path)
