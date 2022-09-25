"""
A file to change a .docx file's tables to have its cell's digits replaced with random digits (no leading zeros, no zeros)
Author: Benjamin Martin
Date: 2022-08-28
"""

import re
import random
import docx
#also need python-docx package installed for docx
from datetime import datetime
import os


def repl_fun(match):
    #random.randint has inclusive boundaries
    return str(random.randint(0, 9))


def randomize_table_digits(table, row_lower_bound, row_upper_bound, col_lower_bound, col_upper_bound, repl_fun,
                               repl_pattern='[0-9]', initialization_flag=False, initialization_str=None):
    """randomizes table of math questions' digits.

    Use regular expression to replace a .docx file table cell's paragraph text digits to other random digits to other
    random digits. If initialization_flag is True, replacement happens after the first instance of initialization_str.

    :param table: docx table
    :param row_lower_bound: int
    :param row_upper_bound: int
    :param col_lower_bound: int
    :param col_upper_bound: int
    :param repl_fun: function for regular expression with argument match
    :param repl_pattern: str of pattern for regular expression
    :param initialization_flag: bool
    :param initialization_str: str
    :return: docx table
    """
    for c in range(col_lower_bound, col_upper_bound + 1):
        for r in range(row_lower_bound, row_upper_bound + 1):
            font_size = table.cell(r, c).paragraphs[0].runs[0].font.size
            s = table.cell(r, c).text

            if initialization_flag:
                initialization_index = s.find(initialization_str)

                if initialization_index == -1:
                    raise ValueError(f'The initilization_str, {initialization_str}, was not found in: {s}. Aborting.')

                prefix = s[:initialization_index + len(initialization_str)]
                question = s[initialization_index + len(initialization_str):]

            else:
                prefix = ''
                question = s

            question_segments = re.split(r'(\d+)', question)

            for ind, seg in enumerate(question_segments):
                replaced_seg = re.sub(repl_pattern, repl_fun, seg)
                if replaced_seg:
                    if replaced_seg[0] == '0':
                        replaced_seg = str(random.randint(1, 9)) + replaced_seg[1:]

                question_segments[ind] = replaced_seg

            replaced_question = ''.join(question_segments)
            table.cell(r, c).text = prefix + replaced_question
            table.cell(r, c).paragraphs[0].runs[0].font.size = font_size

    return table


def add_table_answers(table, row_lower_bound, row_upper_bound, col_lower_bound, col_upper_bound,
                      initialization_flag=False, initialization_str=None):
    """Add answers to of math questions

    :param table: docx table
    :param row_lower_bound: int
    :param row_upper_bound: int
    :param col_lower_bound: int
    :param col_upper_bound: int
    :param initialization_flag: bool
    :param initialization_str: str
    :return: docx table
    """
    operator_translation = {
        '**': ['\^'],
        '*': ['\u00D7', 'x'],
        '/': ['\u00F7']
    }
    for c in range(col_lower_bound, col_upper_bound + 1):
        for r in range(row_lower_bound, row_upper_bound + 1):
            font_size = table.cell(r, c).paragraphs[0].runs[0].font.size
            s = table.cell(r, c).text

            if initialization_flag:
                initialization_index = s.find(initialization_str)

                if initialization_index == -1:
                    raise ValueError(f'The initilization_str, {initialization_str}, was not found in: {s}. Aborting.')

                prefix = s[:initialization_index + len(initialization_str)]
                question = s[initialization_index + len(initialization_str):]

            else:
                prefix = ''
                question = s

            computer_readable_question = question
            for key, values in operator_translation.items():
                for value in values:
                    computer_readable_question = re.sub(value, key, computer_readable_question)
            try:
                answer = eval(computer_readable_question)
                answer_str = ' = '
                if isinstance(answer, int):
                    answer_str += str(answer)

                else:
                    answer_str += format(answer, '.4f')

            except:
                answer_str = ' not evaluable'

            table.cell(r, c).text = prefix + question + answer_str
            table.cell(r, c).paragraphs[0].runs[0].font.size = font_size

    return table


if __name__ == "__main__":
    now = datetime.now().strftime("%m_%d_%Y %H_%M_%S")
    input_file = 'Paragraph Test.docx'
    randomize_test_flag = True
    initialization_flag = True
    make_answer_key_flag = True
    initialization_str = ')'
    randomized_row_lower_bound = 1
    randomized_row_upper_bound = 14
    randomized_col_lower_bound = 1
    randomized_col_upper_bound = 3

    output_test_path = os.path.join(os.getcwd(), 'randomized_docx_tables')

    if not os.path.exists(output_test_path):
        print('Making randomized_docx_tables folder: ', output_test_path)
        os.makedirs(output_test_path)

    if randomize_test_flag:
        output_test_file_path = os.path.join(output_test_path, now + ' ' + input_file)
        test_doc = docx.Document(input_file)

        for i, table in enumerate(test_doc.tables):
            test_doc.tables[i] = randomize_table_digits(table, randomized_row_lower_bound, randomized_row_upper_bound,
                                                    randomized_col_lower_bound, randomized_col_upper_bound, repl_fun,
                                                    initialization_flag=initialization_flag, initialization_str=initialization_str)

        test_doc.save(output_test_file_path)
        print('Outputted randomized test to file: ', output_test_file_path)

    if make_answer_key_flag:
        output_answer_key_file_path = os.path.join(output_test_path, now + ' ANSWER KEY ' + input_file)

        if randomize_test_flag:
            answer_key_doc = docx.Document(output_test_file_path)

        else:
            answer_key_doc = docx.Document(input_file)

        for i, table in enumerate(answer_key_doc.tables):
            answer_key_doc.tables[i] = add_table_answers(table, randomized_row_lower_bound, randomized_row_upper_bound,
                                                randomized_col_lower_bound, randomized_col_upper_bound,
                                                initialization_flag=initialization_flag, initialization_str=initialization_str)

        answer_key_doc.save(output_answer_key_file_path)
        print('Outputted randomized test answer key to file: ', output_answer_key_file_path)